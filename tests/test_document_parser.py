"""document_parser 的单元测试."""
from pathlib import Path

import pytest

from core.tools.document_parser import (
    SUPPORTED_EXTS,
    UnsupportedFileType,
    parse_any,
    parse_docx,
    parse_pdf,
    parse_text,
)


class TestParseText:
    def test_reads_txt(self, tmp_path: Path):
        p = tmp_path / "a.txt"
        p.write_text("hello world", encoding="utf-8")
        assert parse_text(p) == "hello world"

    def test_reads_md(self, tmp_path: Path):
        p = tmp_path / "a.md"
        p.write_text("# title", encoding="utf-8")
        assert parse_text(p) == "# title"

    def test_missing_file(self, tmp_path: Path):
        with pytest.raises(FileNotFoundError):
            parse_text(tmp_path / "nope.txt")


class TestParseAnyDispatch:
    def test_unsupported_ext(self, tmp_path: Path):
        p = tmp_path / "a.xlsx"
        p.write_text("x", encoding="utf-8")
        with pytest.raises(UnsupportedFileType):
            parse_any(p)

    def test_supported_set(self):
        assert {".pdf", ".docx", ".txt", ".md"} <= SUPPORTED_EXTS


class TestParsePdf:
    def test_extracts_text(self, tmp_path: Path):
        # 用 PyMuPDF 现场造一个简单 PDF
        import fitz
        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 72), "Hello PDF World")
        page.insert_text((50, 100), "Second line of text here")
        doc.save(pdf_path)
        doc.close()

        text = parse_pdf(pdf_path)
        assert "Hello PDF World" in text
        assert "Second line" in text

    def test_missing_pdf(self, tmp_path: Path):
        with pytest.raises(FileNotFoundError):
            parse_pdf(tmp_path / "nope.pdf")


class TestParseDocx:
    def test_extracts_paragraphs(self, tmp_path: Path):
        import docx as docx_lib
        docx_path = tmp_path / "test.docx"
        d = docx_lib.Document()
        d.add_paragraph("First paragraph")
        d.add_paragraph("Second paragraph")
        d.save(docx_path)

        text = parse_docx(docx_path)
        assert "First paragraph" in text
        assert "Second paragraph" in text
