"""简历/JD 文件解析工具 - 把 PDF / DOCX / TXT 转成纯文本.

支持格式:
- .pdf  (PyMuPDF)
- .docx (python-docx)
- .txt  / .md (直接读)

不是生产级 OCR (扫描件 PDF 不支持), 假设简历是文本型 PDF.
"""
from __future__ import annotations

import logging
from pathlib import Path

import fitz  # PyMuPDF
import docx

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}


class UnsupportedFileType(ValueError):
    pass


def parse_pdf(path: str | Path) -> str:
    """抽取 PDF 全文, 每页之间用换页符隔开."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"pdf not found: {path}")
    doc = fitz.open(path)
    try:
        parts: list[str] = []
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    finally:
        doc.close()


def parse_docx(path: str | Path) -> str:
    """抽取 docx 全文 (按段落)."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"docx not found: {path}")
    d = docx.Document(path)
    parts: list[str] = []
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    # 表格也补上
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_text(path: str | Path) -> str:
    """读 txt / md 全文."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"text not found: {path}")
    return path.read_text(encoding="utf-8", errors="replace")


def parse_any(path: str | Path) -> str:
    """按后缀分派."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".docx":
        return parse_docx(path)
    if ext in {".txt", ".md"}:
        return parse_text(path)
    raise UnsupportedFileType(f"unsupported: {ext}, supported: {SUPPORTED_EXTS}")
